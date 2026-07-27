"""Enhanced Multi-Modal Document Ingestion Pipeline"""

import logging
import base64
from pathlib import Path
from typing import Dict, List, Tuple, Optional
from io import BytesIO
import zipfile
import warnings
import os
import shutil
import sys

# Suppress PyMuPDF color warnings
warnings.filterwarnings("ignore", message=".*Cannot set gray non-stroke color.*")

# Suppress PyMuPDF stderr messages
import contextlib

@contextlib.contextmanager
def suppress_stderr():
    """Context manager to suppress stderr output"""
    with open(os.devnull, 'w') as devnull:
        old_stderr = sys.stderr
        sys.stderr = devnull
        try:
            yield
        finally:
            sys.stderr = old_stderr

# Core document processing
import pymupdf
from docx import Document as DocxDocument
from PIL import Image
import pandas as pd

# OCR and image processing
import pytesseract
import cv2
import numpy as np

# Configure Tesseract path
if not shutil.which('tesseract'):
    tesseract_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe'
    ]
    for path in tesseract_paths:
        if os.path.exists(path):
            pytesseract.pytesseract.tesseract_cmd = path
            break

# Table extraction
import camelot
import pdfplumber
from tabula import read_pdf

# Suppress pdfminer logging warnings
logging.getLogger("pdfminer").setLevel(logging.ERROR)

logger = logging.getLogger(__name__)


class MultiModalDocumentProcessor:
    """Enhanced document processor for multi-modal RAG system"""
    
    def __init__(self):
        self.supported_formats = {'.pdf', '.docx', '.doc', '.txt'}
        
    def process_documents(self, file_paths: List[Path]) -> List[Dict]:
        """Process multiple documents and extract all content types"""
        results = []
        
        for file_path in file_paths:
            try:
                result = self._process_single_document(file_path)
                results.append(result)
            except Exception as e:
                logger.error(f"Failed to process {file_path}: {e}")
                results.append({
                    'file_path': str(file_path),
                    'error': str(e),
                    'content': {}
                })
                
        return results
    
    def _process_single_document(self, file_path: Path) -> Dict:
        """Process a single document extracting text, tables, and images"""
        ext = file_path.suffix.lower()
        
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported format: {ext}")
            
        content = {
            'text': '',
            'tables': [],
            'images': [],
            'metadata': {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'pages': 0
            }
        }
        
        if ext == '.pdf':
            content = self._process_pdf(file_path)
        elif ext == '.docx':
            content = self._process_docx(file_path)
        elif ext == '.txt':
            content['text'] = file_path.read_text(encoding='utf-8')
            
        return {
            'file_path': str(file_path),
            'content': content
        }
    
    def _process_pdf(self, file_path: Path) -> Dict:
        """Extract text, tables, and images from PDF"""
        content = {
            'text': '',
            'tables': [],
            'images': [],
            'metadata': {
                'file_name': file_path.name,
                'file_size': file_path.stat().st_size,
                'pages': 0
            }
        }
        
        # Text extraction with OCR fallback
        with suppress_stderr():
            doc = pymupdf.open(str(file_path))
            total_pages = len(doc)
        
        content['metadata']['pages'] = total_pages
        logger.info(f"Processing {file_path.name}: {total_pages} pages")
        
        all_text = []
        
        for page_num, page in enumerate(doc, 1):
            # Extract text
            with suppress_stderr():
                page_text = page.get_text().strip()
            
            # OCR fallback for low-text pages
            if len(page_text) < 50:
                page_text = self._ocr_page(page)
                
            all_text.append(f"## Page {page_num}\n{page_text}")
            
            # Extract images
            images = self._extract_pdf_images(page, page_num)
            content['images'].extend(images)
            
        content['text'] = '\n\n'.join(all_text)
        doc.close()
        
        # Extract tables using multiple methods
        content['tables'] = self._extract_pdf_tables(file_path)
        logger.info(f"Extracted {len(content['tables'])} tables, {len(content['images'])} images")
        
        return content
    
    def _process_docx(self, file_path: Path) -> Dict:
        """Extract text, tables, and images from DOCX"""
        content = {
            'text': '',
            'tables': [],
            'images': [],
            'metadata': {'pages': 1}
        }
        
        doc = DocxDocument(file_path)
        
        # Extract text
        paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
        content['text'] = '\n'.join(paragraphs)
        
        # Extract tables
        for i, table in enumerate(doc.tables):
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            
            if table_data:
                content['tables'].append({
                    'table_id': f"table_{i+1}",
                    'data': table_data,
                    'source': 'docx_native'
                })
        
        # Extract images from DOCX
        content['images'] = self._extract_docx_images(file_path)
        
        return content
    
    def _ocr_page(self, page) -> str:
        """Perform OCR on a PDF page"""
        try:
            pix = page.get_pixmap(dpi=300)
            img_data = pix.tobytes("png")
            img = Image.open(BytesIO(img_data))
            
            # Convert to numpy array for OpenCV processing
            img_array = np.array(img)
            
            # Preprocess image for better OCR
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            processed = cv2.threshold(gray, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)[1]
            
            # OCR with Tesseract
            text = pytesseract.image_to_string(processed, config='--psm 6')
            return text.strip()
            
        except Exception as e:
            logger.error(f"OCR failed: {e}")
            return "[OCR failed]"
    
    def _extract_pdf_images(self, page, page_num: int) -> List[Dict]:
        """Extract images from PDF page with chart detection"""
        images = []
        
        try:
            with suppress_stderr():
                image_list = page.get_images()
            
            for img_index, img in enumerate(image_list):
                xref = img[0]
                with suppress_stderr():
                    pix = pymupdf.Pixmap(page.parent, xref)
                
                if pix.n - pix.alpha < 4:
                    img_data = pix.tobytes("png")
                    img_pil = Image.open(BytesIO(img_data))
                    
                    if img_pil.size[0] > 100 and img_pil.size[1] > 100:
                        img_array = np.array(img_pil)
                        
                        # Detect if image is a chart
                        chart_metadata = self._analyze_chart(img_array)
                        
                        # OCR the image
                        ocr_text = pytesseract.image_to_string(img_pil)
                        
                        images.append({
                            'image_id': f"page_{page_num}_img_{img_index+1}",
                            'page': page_num,
                            'size': img_pil.size,
                            'ocr_text': ocr_text.strip(),
                            'format': 'png',
                            'is_chart': chart_metadata['is_chart'],
                            'chart_type': chart_metadata.get('chart_type'),
                            'chart_elements': chart_metadata.get('elements', {})
                        })
                
                pix = None
                
        except Exception as e:
            logger.error(f"Image extraction failed for page {page_num}: {e}")
            
        return images
    
    def _analyze_chart(self, img_array: np.ndarray) -> Dict:
        """Analyze image to detect chart type and extract metadata"""
        try:
            # Convert to grayscale
            if len(img_array.shape) == 3:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = img_array
            
            # Detect edges
            edges = cv2.Canny(gray, 50, 150)
            
            # Detect lines (axes, bars)
            lines = cv2.HoughLinesP(edges, 1, np.pi/180, threshold=100, minLineLength=50, maxLineGap=10)
            
            # Detect contours (pie slices, bars)
            contours, _ = cv2.findContours(edges, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            # Chart detection heuristics
            is_chart = False
            chart_type = None
            elements = {}
            
            if lines is not None and len(lines) > 10:
                is_chart = True
                
                # Analyze line orientations
                horizontal_lines = 0
                vertical_lines = 0
                
                for line in lines:
                    x1, y1, x2, y2 = line[0]
                    angle = np.abs(np.arctan2(y2 - y1, x2 - x1) * 180 / np.pi)
                    
                    if angle < 10 or angle > 170:
                        horizontal_lines += 1
                    elif 80 < angle < 100:
                        vertical_lines += 1
                
                elements['horizontal_lines'] = horizontal_lines
                elements['vertical_lines'] = vertical_lines
                elements['total_lines'] = len(lines)
                
                # Classify chart type
                if horizontal_lines > 5 and vertical_lines > 5:
                    # Grid pattern suggests bar/line chart
                    if len(contours) > 3:
                        # Multiple rectangular contours = bar chart
                        rect_contours = sum(1 for c in contours if len(cv2.approxPolyDP(c, 0.04 * cv2.arcLength(c, True), True)) == 4)
                        if rect_contours > 3:
                            chart_type = 'bar_chart'
                            elements['bars_detected'] = rect_contours
                        else:
                            chart_type = 'line_chart'
                    else:
                        chart_type = 'line_chart'
                elif len(contours) > 0:
                    # Check for circular contours (pie chart)
                    for contour in contours:
                        area = cv2.contourArea(contour)
                        perimeter = cv2.arcLength(contour, True)
                        if perimeter > 0:
                            circularity = 4 * np.pi * area / (perimeter ** 2)
                            if circularity > 0.7:
                                chart_type = 'pie_chart'
                                elements['circular_elements'] = len([c for c in contours if cv2.contourArea(c) > 100])
                                break
            
            # Color analysis for chart detection
            if len(img_array.shape) == 3:
                unique_colors = len(np.unique(img_array.reshape(-1, img_array.shape[2]), axis=0))
                elements['unique_colors'] = unique_colors
                
                # Charts typically have distinct color regions
                if unique_colors > 10 and unique_colors < 50:
                    is_chart = True
                    if chart_type is None:
                        chart_type = 'unknown_chart'
            
            return {
                'is_chart': is_chart,
                'chart_type': chart_type,
                'elements': elements,
                'confidence': self._calculate_chart_confidence(is_chart, chart_type, elements)
            }
            
        except Exception as e:
            logger.debug(f"Chart analysis failed: {e}")
            return {'is_chart': False, 'chart_type': None, 'elements': {}}
    
    def _calculate_chart_confidence(self, is_chart: bool, chart_type: Optional[str], elements: Dict) -> float:
        """Calculate confidence score for chart detection"""
        if not is_chart:
            return 0.0
        
        confidence = 0.5  # Base confidence
        
        # Boost confidence based on detected elements
        if elements.get('total_lines', 0) > 20:
            confidence += 0.2
        if elements.get('horizontal_lines', 0) > 5 and elements.get('vertical_lines', 0) > 5:
            confidence += 0.15
        if chart_type in ['bar_chart', 'line_chart', 'pie_chart']:
            confidence += 0.15
        
        return min(confidence, 1.0)
    
    def _extract_docx_images(self, file_path: Path) -> List[Dict]:
        """Extract images from DOCX file with chart detection"""
        images = []
        
        try:
            with zipfile.ZipFile(file_path, 'r') as docx_zip:
                image_files = [f for f in docx_zip.namelist() 
                             if f.startswith("word/media/") and 
                             f.lower().endswith(('.png', '.jpg', '.jpeg'))]
                
                for i, image_name in enumerate(image_files):
                    with docx_zip.open(image_name) as image_file:
                        img_data = image_file.read()
                        img = Image.open(BytesIO(img_data))
                        
                        if img.size[0] > 100 and img.size[1] > 100:
                            img_array = np.array(img)
                            
                            # Detect if image is a chart
                            chart_metadata = self._analyze_chart(img_array)
                            
                            # OCR the image
                            ocr_text = pytesseract.image_to_string(img)
                            
                            images.append({
                                'image_id': f"docx_img_{i+1}",
                                'filename': image_name,
                                'size': img.size,
                                'ocr_text': ocr_text.strip(),
                                'format': image_name.split('.')[-1].lower(),
                                'is_chart': chart_metadata['is_chart'],
                                'chart_type': chart_metadata.get('chart_type'),
                                'chart_elements': chart_metadata.get('elements', {})
                            })
                            
        except Exception as e:
            logger.error(f"DOCX image extraction failed: {e}")
            
        return images
    
    def _extract_pdf_tables(self, file_path: Path) -> List[Dict]:
        """Extract tables from PDF using multiple methods"""
        tables = []
        
        # Method 1: Camelot (for well-formatted tables)
        try:
            with suppress_stderr():
                camelot_tables = camelot.read_pdf(str(file_path), pages='all')
            for i, table in enumerate(camelot_tables):
                if not table.df.empty:
                    tables.append({
                        'table_id': f"camelot_table_{i+1}",
                        'page': table.page,
                        'data': table.df.values.tolist(),
                        'headers': table.df.columns.tolist(),
                        'source': 'camelot',
                        'accuracy': table.accuracy
                    })
        except Exception as e:
            logger.debug(f"Camelot table extraction failed: {e}")
        
        # Method 2: Tabula (for stream tables)
        try:
            with suppress_stderr():
                tabula_tables = read_pdf(str(file_path), pages='all', multiple_tables=True)
            for i, df in enumerate(tabula_tables):
                if not df.empty:
                    tables.append({
                        'table_id': f"tabula_table_{i+1}",
                        'data': df.values.tolist(),
                        'headers': df.columns.tolist(),
                        'source': 'tabula'
                    })
        except Exception as e:
            logger.debug(f"Tabula table extraction failed: {e}")
        
        # Method 3: pdfplumber (fallback)
        try:
            with suppress_stderr():
                with pdfplumber.open(file_path) as pdf:
                    for page_num, page in enumerate(pdf.pages, 1):
                        page_tables = page.extract_tables()
                        for i, table in enumerate(page_tables):
                            if table and len(table) > 1:
                                tables.append({
                                    'table_id': f"pdfplumber_table_p{page_num}_{i+1}",
                                    'page': page_num,
                                    'data': table,
                                    'source': 'pdfplumber'
                                })
        except Exception as e:
            logger.debug(f"pdfplumber table extraction failed: {e}")
        
        return tables


class DocumentIngestionPipeline:
    """Main pipeline for document ingestion"""
    
    def __init__(self):
        self.processor = MultiModalDocumentProcessor()
    
    def ingest_documents(self, input_path: Path) -> List[Dict]:
        """Ingest documents from file or directory"""
        input_path = Path(input_path)
        if input_path.is_file():
            files = [input_path]
        elif input_path.is_dir():
            files = []
            for ext in ['.pdf', '.docx', '.doc', '.txt']:
                files.extend(input_path.glob(f'*{ext}'))
        else:
            raise ValueError(f"Invalid path: {input_path}")
        
        logger.info(f"Processing {len(files)} documents")
        return self.processor.process_documents(files)
    
    def get_content_summary(self, results: List[Dict]) -> Dict:
        """Get summary of extracted content"""
        summary = {
            'total_documents': len(results),
            'successful': 0,
            'failed': 0,
            'total_text_length': 0,
            'total_tables': 0,
            'total_images': 0
        }
        
        for result in results:
            if 'error' in result:
                summary['failed'] += 1
            else:
                summary['successful'] += 1
                content = result['content']
                summary['total_text_length'] += len(content.get('text', ''))
                summary['total_tables'] += len(content.get('tables', []))
                summary['total_images'] += len(content.get('images', []))
        
        return summary